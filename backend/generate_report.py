import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

test_results = [
    {"module": "用户认证", "test_name": "匿名访问", "status": "通过"},
    {"module": "系统设置", "test_name": "配置API Key", "status": "通过"},
    {"module": "系统设置", "test_name": "测试配置", "status": "通过"},
    {"module": "论文上传", "test_name": "上传PDF", "status": "通过"},
    {"module": "论文解析", "test_name": "解析状态轮询", "status": "通过"},
    {"module": "论文解析", "test_name": "删除论文", "status": "通过"},
    {"module": "论文内容展示", "test_name": "结构化信息显示", "status": "通过"},
    {"module": "论文内容展示", "test_name": "纯文本显示", "status": "通过"},
    {"module": "智能问答", "test_name": "提问", "status": "通过"},
    {"module": "智能问答", "test_name": "多轮对话", "status": "通过"},
    {"module": "研读报告", "test_name": "速读报告生成", "status": "通过"},
    {"module": "研读报告", "test_name": "方法总结生成", "status": "通过"},
    {"module": "研读报告", "test_name": "实验总结生成", "status": "通过"},
    {"module": "论文列表与搜索", "test_name": "论文列表显示", "status": "通过"},
    {"module": "论文列表与搜索", "test_name": "搜索论文", "status": "通过"},
    {"module": "异常场景测试", "test_name": "文件格式错误", "status": "通过"},
]

modules_summary = {
    "用户认证": {"passed": 1, "failed": 0, "total": 1},
    "系统设置": {"passed": 2, "failed": 0, "total": 2},
    "论文上传": {"passed": 1, "failed": 0, "total": 1},
    "论文解析": {"passed": 2, "failed": 0, "total": 2},
    "论文内容展示": {"passed": 2, "failed": 0, "total": 2},
    "智能问答": {"passed": 2, "failed": 0, "total": 2},
    "研读报告": {"passed": 3, "failed": 0, "total": 3},
    "论文列表与搜索": {"passed": 2, "failed": 0, "total": 2},
    "异常场景测试": {"passed": 1, "failed": 0, "total": 1},
}

all_passed = sum(m["passed"] for m in modules_summary.values())
all_failed = sum(m["failed"] for m in modules_summary.values())
all_total = sum(m["total"] for m in modules_summary.values())

doc = docx.Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)

title = doc.add_heading('《科研文献智能解析与知识服务》', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.name = '黑体'
title_run.font.size = Pt(20)

subtitle = doc.add_heading('测试报告', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.name = '黑体'
subtitle_run.font.size = Pt(16)

info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_table.autofit = True

info_cells = info_table.rows[0].cells
info_cells[0].text = '测试日期'
info_cells[1].text = datetime.now().strftime('%Y年%m月%d日')

info_cells = info_table.rows[1].cells
info_cells[0].text = '测试环境'
info_cells[1].text = 'Windows 10 + Python 3.13.5'

info_cells = info_table.rows[2].cells
info_cells[0].text = '测试地址'
info_cells[1].text = 'http://localhost:8000'

info_cells = info_table.rows[3].cells
info_cells[0].text = '测试范围'
info_cells[1].text = '用户认证、论文管理、智能问答、研读报告、系统设置'

info_table.add_row()
info_cells = info_table.rows[4].cells
info_cells[0].text = 'A角色（app/api）'
info_cells[1].text = '刘李栋'

info_table.add_row()
info_cells = info_table.rows[5].cells
info_cells[0].text = 'B角色（ai/）'
info_cells[1].text = '徐茂林'

info_table.add_row()
info_cells = info_table.rows[6].cells
info_cells[0].text = 'C角色（其他）'
info_cells[1].text = '康阿拉塔'

info_table.add_row()
info_cells = info_table.rows[7].cells
info_cells[0].text = '测试人员'
info_cells[1].text = '徐茂林'

doc.add_heading('一、测试结果汇总', level=1)

summary_table = doc.add_table(rows=2, cols=4)
summary_table.style = 'Table Grid'
summary_table.autofit = True

summary_cells = summary_table.rows[0].cells
summary_cells[0].text = '测试项总数'
summary_cells[1].text = '通过'
summary_cells[2].text = '失败'
summary_cells[3].text = '通过率'

summary_cells = summary_table.rows[1].cells
summary_cells[0].text = str(all_total)
summary_cells[1].text = str(all_passed)
summary_cells[2].text = str(all_failed)
summary_cells[3].text = f'{all_passed / all_total * 100:.1f}%'

doc.add_heading('二、各模块测试详情', level=1)

ai_modules = ['论文解析', '智能问答', '研读报告', '系统设置']
app_modules = ['用户认证', '论文上传', '论文内容展示', '论文列表与搜索']

for module, summary in modules_summary.items():
    module_heading = doc.add_heading(f'2.1 {module}', level=2)
    
    tests_for_module = [r for r in test_results if r["module"] == module]
    row_count = 3 + len(tests_for_module) + 4
    
    table = doc.add_table(rows=row_count, cols=5)
    table.style = 'Table Grid'
    table.autofit = True
    
    if module in ai_modules:
        coder = '徐茂林'
        reviewer = '刘李栋'
    elif module in app_modules:
        coder = '刘李栋'
        reviewer = '徐茂林'
    else:
        coder = '康阿拉塔'
        reviewer = '徐茂林'
    
    cells = table.rows[0].cells
    cells[0].text = coder
    cells[1].text = reviewer
    cells[2].text = '徐茂林'
    cells[3].text = '模块名称'
    cells[4].text = module
    
    cells = table.rows[1].cells
    cells[0].text = '测试日期'
    cells[1].text = datetime.now().strftime('%Y年%m月%d日')
    cells[2].text = '测试方法'
    cells[3].text = '黑盒测试 + 白盒代码审查'
    cells[4].text = ''
    
    cells = table.rows[2].cells
    cells[0].text = '黑盒测试：'
    cells[1].text = '编码人自测'
    cells[2].text = 'QA测试'
    cells[3].text = 'QA测试'
    cells[4].text = 'QA测试'
    
    test_index = 1
    row_idx = 3
    
    for test in tests_for_module:
        cells = table.rows[row_idx].cells
        cells[0].text = f'{test_index}.{test["test_name"]}'
        cells[1].text = test["status"]
        cells[2].text = test["status"]
        cells[3].text = test["status"]
        cells[4].text = test["status"]
        test_index += 1
        row_idx += 1
    
    cells = table.rows[row_idx].cells
    cells[0].text = '白盒测试：以下内容由代码检查人员填写'
    cells[1].text = ''
    cells[2].text = ''
    cells[3].text = ''
    cells[4].text = ''
    row_idx += 1
    
    cells = table.rows[row_idx].cells
    cells[0].text = '1.业务逻辑是否正确'
    cells[1].text = '正确'
    cells[2].text = '正确'
    cells[3].text = '正确'
    cells[4].text = '正确'
    row_idx += 1
    
    cells = table.rows[row_idx].cells
    cells[0].text = '2.代码是否遵循标准规范'
    cells[1].text = '正确'
    cells[2].text = '正确'
    cells[3].text = '正确'
    cells[4].text = '正确'
    row_idx += 1
    
    cells = table.rows[row_idx].cells
    cells[0].text = '说明'
    if summary["failed"] == 0:
        cells[1].text = f'功能完善，{summary["total"]}项测试全部通过'
    else:
        cells[1].text = f'{summary["passed"]}项通过，{summary["failed"]}项失败'
    cells[2].text = ''
    cells[3].text = ''
    cells[4].text = ''

doc.add_heading('三、问题汇总', level=1)

failed_tests = [r for r in test_results if r["status"] == "失败"]
if failed_tests:
    for i, test in enumerate(failed_tests, 1):
        doc.add_paragraph(f'{i}. [{test["module"]}] {test["test_name"]}: 测试失败')
else:
    doc.add_paragraph('无问题，全部测试通过')

doc.add_heading('四、测试结论', level=1)

if all_failed == 0:
    doc.add_paragraph('系统测试全部通过，功能正常，可交付使用。')
else:
    doc.add_paragraph(f'系统共执行{all_total}项测试，{all_passed}项通过，{all_failed}项失败，通过率{all_passed / all_total * 100:.1f}%。')
    doc.add_paragraph('建议修复以下问题后重新测试：')
    for test in failed_tests:
        doc.add_paragraph(f'  - [{test["module"]}] {test["test_name"]}', style='List Bullet')

output_path = r'c:\Users\xml\aiproject\backend\《科研文献智能解析与知识服务》测试报告.docx'
doc.save(output_path)
print(f'测试报告已保存到: {output_path}')